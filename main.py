import os
import telebot
from telebot import types
from pymongo import MongoClient
from docx import Document
from docx.shared import Inches

token = os.environ.get("SOLARIS_TOKEN")
bot = telebot.TeleBot(token)

client = MongoClient('mongodb://localhost:27017/')
db = client['RatingSolaris']
collection = db['Ratings']

data = {}
is_start = {}

debug = True

YULIA = "razumova1983"
if debug:
    YULIA = "pi31415926535qwerty"


@bot.message_handler(commands=['start'])
def start(message):
    if message.from_user.username == YULIA:
        bot.send_message(message.chat.id,
                         "Здравствуйте, Юлия!\nПожалуйста, пришлите команду /get для получения данных.")
        return
    if is_start.get(message.chat.id):
        bot.send_message(message.chat.id, "Вначале завершите Предыдщуий отзыв")
    else:
        is_start[message.chat.id] = True
        keyboard = types.InlineKeyboardMarkup()
        key_start = types.InlineKeyboardButton(text="Начать", callback_data="start")
        keyboard.add(key_start)
        text = """Здравствуйте! 🌟

Благодарим вас за выбор программы от Соляриса! Нам важно ваше мнение, и мы хотели бы услышать ваш отзыв.

Обратная связь поможет нам улучшить качество обучения и сделать курсы еще лучше. 

Пожалуйста, уделите несколько минут, чтобы ответить на несколько вопросов и поделиться своими впечатлениями!"""
        bot.send_message(message.from_user.id, text, reply_markup=keyboard)


@bot.message_handler(commands=['get'])
def get(message):
    if message.from_user.username == YULIA:
        dictionary = list(collection.find())
        if not dictionary:
            bot.send_message(message.chat.id, "Ещё нет ни одного отзыва!")
            return
        bot.send_message(message.chat.id, "Принято, высылаем файл с данными...")
        document = Document()
        document.add_heading("Отзывы Солярис", level=0)
        for i in list(collection.find()):
            p = document.add_paragraph("\nФИО: ")
            p.add_run(i['name']).bold = True

            p.add_run("\nВозраст: ")
            p.add_run(str(i['age'])).bold = True

            p.add_run("\nКласс: ")
            p.add_run(str(i['class'])).bold = True

            p.add_run("\nКурс: ")
            p.add_run(str(i['direction'])).bold = True

            p.add_run("\nПроживали ли в общежитии?: ")
            p.add_run(str(i['q1'])).bold = True

            p.add_run("\nУстроили ли условия проживания?: ")
            p.add_run(str(i['q2'])).bold = True

            if "q2_feedback" in i:
                p.add_run("\nЧто не устроило: ")
                p.add_run(str(i['q2_feedback'])).bold = True

            p.add_run("\nБыли ли полезны лекции?: ")
            p.add_run(str(i['q3'])).bold = True

            p.add_run("\nСтали бы вы рекомендовать наши курсы друзьям/знакомым?: ")
            p.add_run(str(i['q4'])).bold = True

            p.add_run("\nПосетили бы вы другие курсы/смены/интенсивы в Солярисе еще раз?: ")
            p.add_run(str(i['q4'])).bold = True

            p.add_run("\nРазвёрнутый отзыв: ")
            p.add_run(str(i['detailed_feedback'])).bold = True

            p.add_run("\nОбщая оценка: ")
            p.add_run(str(i['rating'])).bold = True

            document.add_paragraph("\n" + "_" * 100)
            collection.delete_one({"name": i['name']})
        document.save("reviews.docx")
        bot.send_document(message.chat.id, open("reviews.docx", "rb"))
        os.remove("reviews.docx")


@bot.callback_query_handler(func=lambda call: call.data == "start")
def callback_worker(call):
    message = call.message
    bot.edit_message_text(chat_id=message.chat.id, message_id=message.message_id, text="Введите ваше ФИО.")
    bot.register_next_step_handler(message, get_name)


def get_name(message):
    data[message.chat.id] = {
        "user_id": message.chat.id,
        "username": message.from_user.username,
        "name": message.text,
        "age": 0,
        "direction": "",
        "q1": "",
        "q2": "",
        "q3": "",
        "q4": "",
        "q5": "",
        "class": "",
        "detailed_feedback": "",
        "rating": 0
    }
    bot.send_message(chat_id=message.chat.id, text="В каком классе вы обучаетесь?")
    bot.register_next_step_handler(message, get_class)


def get_class(message):
    data[message.chat.id]["class"] = message.text
    bot.send_message(chat_id=message.chat.id, text="Введите ваш возраст.")
    bot.register_next_step_handler(message, get_age)


def get_courses():
    with open('courses.txt', encoding='utf8') as f:
        courses = f.read().strip().split('\n')
    return courses


def get_age(message):
    data[message.chat.id]["age"] = message.text
    text = "Напишите название курса/интенсива/смены который вы прошли:"
    bot.send_message(chat_id=message.chat.id, text=text)
    bot.register_next_step_handler(message, get_course)


def get_course(message):
    data[message.chat.id]["direction"] = message.text
    keyboard = types.InlineKeyboardMarkup()
    yes = types.InlineKeyboardButton(text="Да", callback_data="q1_yes")
    keyboard.add(yes)
    no = types.InlineKeyboardButton(text="Нет", callback_data="q1_no")
    keyboard.add(no)

    bot.send_message(chat_id=message.chat.id,
                     text="Проживали ли вы в общежитие во время проведения курса?",
                     reply_markup=keyboard)


@bot.callback_query_handler(func=lambda call: call.data.startswith("q1"))
def callback_q1(call):
    message = call.message
    if call.data == "q1_yes":
        data[message.chat.id]["q1"] = "Да"
        keyboard = types.InlineKeyboardMarkup()
        yes = types.InlineKeyboardButton(text="Да", callback_data="q2_yes")
        keyboard.add(yes)
        no = types.InlineKeyboardButton(text="Нет", callback_data="q2_no")
        keyboard.add(no)
        bot.edit_message_text(chat_id=message.chat.id, message_id=message.message_id,
                              text="Устраивали ли вас условия проживания?", reply_markup=keyboard)
    else:
        data[message.chat.id]["q1"] = "Нет"
        ask_q3(message)


def ask_q3(message, edit=True):
    keyboard = types.InlineKeyboardMarkup()
    yes = types.InlineKeyboardButton(text="Да", callback_data="q3_yes")
    keyboard.add(yes)
    no = types.InlineKeyboardButton(text="Нет", callback_data="q3_no")
    keyboard.add(no)
    if not edit:
        bot.send_message(chat_id=message.chat.id, text="Были ли вам полезны лекции?",
                         reply_markup=keyboard)
    else:
        bot.edit_message_text(chat_id=message.chat.id, message_id=message.message_id,
                              text="Были ли вам полезны лекции?", reply_markup=keyboard)


@bot.callback_query_handler(func=lambda call: call.data.startswith("q2"))
def callback_q2(call):
    message = call.message
    if call.data == "q2_yes":
        data[message.chat.id]["q2"] = "Да"
        ask_q3(message)
    else:
        data[message.chat.id]["q2"] = "Нет"
        bot.edit_message_text(chat_id=message.chat.id, message_id=message.message_id,
                              text="Напишите конкретно что именно вам не понравилось")
        bot.register_next_step_handler(message, get_dop_info)


def get_dop_info(message):
    data[message.chat.id]["q2_feedback"] = message.text
    ask_q3(message, edit=False)


@bot.callback_query_handler(func=lambda call: call.data.startswith("q3"))
def callback_q3(call):
    message = call.message
    if call.data == "q3_yes":
        data[message.chat.id]["q3"] = "Да"
    else:
        data[message.chat.id]["q3"] = "Нет"
    keyboard = types.InlineKeyboardMarkup()
    yes = types.InlineKeyboardButton(text="Да", callback_data="q4_yes")
    keyboard.add(yes)
    no = types.InlineKeyboardButton(text="Нет", callback_data="q4_no")
    keyboard.add(no)
    bot.edit_message_text(chat_id=message.chat.id, message_id=message.message_id,
                          text="Стали бы вы рекомендовать наши курсы друзьям/знакомым?",
                          reply_markup=keyboard)


@bot.callback_query_handler(func=lambda call: call.data.startswith("q4"))
def callback_q4(call):
    message = call.message
    if call.data == "q4_yes":
        data[message.chat.id]["q4"] = "Да"
    else:
        data[message.chat.id]["q4"] = "Нет"
    keyboard = types.InlineKeyboardMarkup()
    yes = types.InlineKeyboardButton(text="Да", callback_data="q5_yes")
    keyboard.add(yes)
    no = types.InlineKeyboardButton(text="Нет", callback_data="q5_no")
    keyboard.add(no)
    bot.edit_message_text(chat_id=message.chat.id, message_id=message.message_id,
                          text="Посетили бы вы другие курсы/смены/интенсивы в Солярисе еще раз?",
                          reply_markup=keyboard)


@bot.callback_query_handler(func=lambda call: call.data.startswith("q5"))
def callback_q5(call):
    message = call.message
    if call.data == "q5_yes":
        data[message.chat.id]["q5"] = "Да"
    else:
        data[message.chat.id]["q5"] = "Нет"
    text = """Развернутый отзыв. 🌠

Напишите про ваши впечатления от курса/смены/интенсива. Расскажите, что вам понравилось, а что бы вы улучшили или изменили."""
    bot.edit_message_text(chat_id=message.chat.id, text=text, message_id=message.message_id)
    bot.register_next_step_handler(message, get_info)


def get_info(message):
    data[message.chat.id]["detailed_feedback"] = message.text
    keyboard = types.InlineKeyboardMarkup()
    buttons = []
    for i in range(1, 6):
        button = types.InlineKeyboardButton(text=str(i), callback_data=f"q7_{i}")
        buttons.append(button)
    keyboard.row(*buttons)
    bot.send_message(chat_id=message.chat.id, text="Оцените в целом мероприятие от 1 до 5",
                     reply_markup=keyboard)


@bot.callback_query_handler(func=lambda call: call.data.startswith("q7"))
def callback_q7(call):
    message = call.message
    data[message.chat.id]["rating"] = int(call.data[-1])

    user_data = data[message.chat.id]
    collection.insert_one(user_data)
    is_start[message.chat.id] = False

    text = """Спасибо за ваш отзыв! 💫

Мы ценим ваше мнение и время, которое вы потратили на опрос. Обратная связь очень важна для улучшения нашей работы. 

Желаем вам удачи в учебе и надеемся увидеть вас снова!\nЧтобы оставить ещё один отзыв отправьте команду /start"""
    bot.edit_message_text(chat_id=message.chat.id, message_id=message.message_id,
                          text=text)


bot.polling(none_stop=True, interval=0)
